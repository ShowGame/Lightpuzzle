#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
将 小咪的光学迷宫游戏软件说明.md 转为 Word 文档，并嵌入备案截图。
用法: python tools/build-optical-copyright-doc.py
"""

from __future__ import annotations

import re
import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
PROJECT_ROOT = Path(__file__).resolve().parent.parent
MD_PATH = PROJECT_ROOT / '软著申请材料' / '小咪的光学迷宫游戏软件说明.md'
OUT_PATH = PROJECT_ROOT / '软著申请材料' / '小咪的光学迷宫游戏软件说明.docx'
IMAGE_WIDTH = Cm(14)


def set_run_font(run, name: str = '宋体', size_pt: float = 12, bold: bool = False):
    run.font.name = name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)


def set_paragraph_format(paragraph, first_line_indent_pt: float = 0):
    fmt = paragraph.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    fmt.line_spacing = 1.25
    fmt.space_after = Pt(6)
    if first_line_indent_pt:
        fmt.first_line_indent = Pt(first_line_indent_pt)


def add_styled_paragraph(doc: Document, text: str, style: str | None = None, bold: bool = False):
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    set_run_font(run, bold=bold)
    set_paragraph_format(p, first_line_indent_pt=24 if not style else 0)
    return p


def add_rich_paragraph(doc: Document, text: str):
    """支持 **粗体** 的简单行内格式。"""
    p = doc.add_paragraph()
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            set_run_font(run, bold=True)
        elif part:
            run = p.add_run(part)
            set_run_font(run)
    set_paragraph_format(p, first_line_indent_pt=24)
    return p


def add_code_block(doc: Document, lines: list[str]):
    p = doc.add_paragraph()
    run = p.add_run('\n'.join(lines))
    set_run_font(run, name='Consolas', size_pt=9)
    set_paragraph_format(p)
    p.paragraph_format.left_indent = Cm(0.8)
    return p


def add_table_from_md(doc: Document, rows: list[list[str]]):
    if len(rows) < 2:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = 'Table Grid'
    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(cell_text.strip())
            set_run_font(run, size_pt=10, bold=(r_idx == 0))
    doc.add_paragraph()


def resolve_image_path(raw: str, md_dir: Path) -> Path | None:
    candidate = (md_dir / raw).resolve()
    if candidate.is_file():
        return candidate
    alt = PROJECT_ROOT / raw.lstrip('./')
    if alt.is_file():
        return alt
    return None


def refresh_git_table(md_text: str) -> str:
    """若 git 可用，刷新「二、Git 提交记录」表格。"""
    try:
        result = subprocess.run(
            ['git', 'log', '--format=%h|%ad|%an|%s', '--date=short'],
            cwd=PROJECT_ROOT,
            capture_output=True,
            text=True,
            encoding='utf-8',
            check=True,
        )
    except (subprocess.CalledProcessError, FileNotFoundError):
        return md_text

    lines = [ln for ln in result.stdout.strip().splitlines() if ln.strip()]
    if not lines:
        return md_text

    header = '| 序号 | 短哈希 | 日期 | 作者 | 说明 |\n|------|--------|------|------|------|\n'
    body = []
    for i, ln in enumerate(lines, 1):
        h, date, author, msg = ln.split('|', 3)
        body.append(f'| {i} | `{h}` | {date} | {author} | {msg} |')
    new_table = header + '\n'.join(body)

    pattern = r'(## 二、Git 提交记录\s*\n\n)([\s\S]*?)(\n---\s*\n\n## 三、游戏截图说明)'
    if re.search(pattern, md_text):
        return re.sub(pattern, rf'\1{new_table}\3', md_text, count=1)
    return md_text


def parse_markdown_to_docx(md_path: Path, out_path: Path):
    md_text = md_path.read_text(encoding='utf-8')
    md_text = refresh_git_table(md_text)
    md_dir = md_path.parent

    doc = Document()
    # 默认正文样式
    normal = doc.styles['Normal']
    normal.font.name = '宋体'
    normal.font.size = Pt(12)
    normal._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    lines = md_text.splitlines()
    i = 0
    table_buffer: list[list[str]] = []
    code_buffer: list[str] | None = None

    def flush_table():
        nonlocal table_buffer
        if table_buffer:
            add_table_from_md(doc, table_buffer)
            table_buffer = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if code_buffer is not None:
            if stripped.startswith('```'):
                add_code_block(doc, code_buffer)
                code_buffer = None
            else:
                code_buffer.append(line)
            i += 1
            continue

        if stripped.startswith('```'):
            flush_table()
            code_buffer = []
            i += 1
            continue

        if stripped.startswith('|'):
            cells = [c.strip() for c in stripped.strip('|').split('|')]
            if cells and not all(re.match(r'^[-:]+$', c) for c in cells):
                table_buffer.append(cells)
            i += 1
            continue

        flush_table()

        if not stripped or stripped == '---':
            i += 1
            continue

        img_match = re.match(r'!\[([^\]]*)\]\(([^)]+)\)', stripped)
        if img_match:
            alt, rel = img_match.groups()
            img_path = resolve_image_path(rel, md_dir)
            if img_path:
                doc.add_picture(str(img_path), width=IMAGE_WIDTH)
                cap = doc.add_paragraph()
                run = cap.add_run(alt or img_path.stem)
                set_run_font(run, size_pt=10)
                cap.alignment = 1  # center
            else:
                add_styled_paragraph(doc, f'[图片缺失: {rel}]')
            i += 1
            continue

        if stripped.startswith('# '):
            doc.add_heading(stripped[2:].strip(), level=0)
            i += 1
            continue
        if stripped.startswith('## '):
            doc.add_heading(stripped[3:].strip(), level=1)
            i += 1
            continue
        if stripped.startswith('### '):
            doc.add_heading(stripped[4:].strip(), level=2)
            i += 1
            continue
        if stripped.startswith('#### '):
            doc.add_heading(stripped[5:].strip(), level=3)
            i += 1
            continue
        if stripped.startswith('##### '):
            doc.add_heading(stripped[6:].strip(), level=4)
            i += 1
            continue

        if stripped.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            text = stripped[2:]
            if '**' in text:
                parts = re.split(r'(\*\*[^*]+\*\*)', text)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = p.add_run(part[2:-2])
                        set_run_font(run, bold=True)
                    elif part:
                        run = p.add_run(part)
                        set_run_font(run)
            else:
                run = p.add_run(text)
                set_run_font(run)
            set_paragraph_format(p)
            i += 1
            continue

        if re.match(r'^\d+\.\s', stripped):
            p = doc.add_paragraph(style='List Number')
            run = p.add_run(re.sub(r'^\d+\.\s*', '', stripped))
            set_run_font(run)
            set_paragraph_format(p)
            i += 1
            continue

        if '**' in stripped:
            add_rich_paragraph(doc, stripped)
        else:
            add_styled_paragraph(doc, stripped)
        i += 1

    flush_table()

    out_path.parent.mkdir(parents=True, exist_ok=True)
    try:
        doc.save(str(out_path))
    except PermissionError:
        fallback = out_path.with_name(out_path.stem + '_生成.docx')
        doc.save(str(fallback))
        print(f'原文件被占用，已写入: {fallback}')
        print('请关闭 Word 后重新运行脚本，或手动将生成文件替换原文件。')
        return fallback
    return out_path


def main():
    if not MD_PATH.is_file():
        print(f'未找到 Markdown: {MD_PATH}', file=sys.stderr)
        sys.exit(1)
    out = parse_markdown_to_docx(MD_PATH, OUT_PATH)
    print(f'已生成: {out}')


if __name__ == '__main__':
    main()
